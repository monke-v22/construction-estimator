"""
tests/test_phase1_parsers.py
Phase 1 tests for all document parsers and Agent #1 dispatcher.
Uses synthetically generated test files — no real API calls needed.
"""

import sys
import json
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest


# ─────────────────────────────────────────────
# Helpers — create synthetic test files
# ─────────────────────────────────────────────

def make_test_pdf(tmp_path: Path, text: str = None) -> Path:
    """Create a minimal text-based PDF using reportlab or fpdf2."""
    pdf_path = tmp_path / "test_boq.pdf"
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        c.setFont("Helvetica", 10)
        lines = (text or SAMPLE_BOQ_TEXT).split("\n")
        y = 800
        for line in lines[:50]:
            c.drawString(50, y, line[:100])
            y -= 15
            if y < 50:
                c.showPage()
                y = 800
        c.save()
    except ImportError:
        # Fallback: write minimal valid PDF bytes
        pdf_content = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]
/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 44>>
stream
BT /F1 12 Tf 72 720 Td (BOQ Item 1 - Tiling 100 m2) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000370 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
441
%%EOF"""
        pdf_path.write_bytes(pdf_content)
    return pdf_path


def make_test_docx(tmp_path: Path) -> Path:
    """Create a .docx with paragraphs and a table."""
    from docx import Document
    path = tmp_path / "test_spec.docx"
    doc = Document()
    doc.add_heading("FIT-OUT SPECIFICATION", level=1)
    doc.add_paragraph("Project: Office Fit-Out — Floor 3, Tower A")
    doc.add_paragraph("Total Area: 450 m2")
    doc.add_paragraph("Grade: Premium")

    table = doc.add_table(rows=4, cols=4)
    headers = ["Item", "Description", "Unit", "Qty"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["1.01", "Raised access flooring", "m2", "350"],
        ["1.02", "Ceiling grid system", "m2", "350"],
        ["1.03", "Glass partitions", "lm", "45"],
    ]
    for r_idx, row_data in enumerate(data, 1):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = val

    doc.save(str(path))
    return path


def make_test_excel(tmp_path: Path) -> Path:
    """Create a .xlsx with a BOQ sheet."""
    import pandas as pd
    path = tmp_path / "test_boq.xlsx"
    df = pd.DataFrame({
        "Item No": ["1.01", "1.02", "1.03", "2.01", "2.02"],
        "Description": [
            "Ceramic floor tiles 600x600",
            "Gypsum board ceiling",
            "Internal paint (2 coats)",
            "Split A/C units (2.5 ton)",
            "Electrical DB board",
        ],
        "Trade": ["finishing", "finishing", "finishing", "mep", "mep"],
        "Unit": ["m2", "m2", "m2", "nr", "nr"],
        "Qty": [350, 300, 800, 12, 3],
        "Rate": [85, 65, 18, 1200, 2500],
        "Amount": [29750, 19500, 14400, 14400, 7500],
    })
    df.to_excel(str(path), index=False, sheet_name="BOQ")
    return path


def make_test_csv(tmp_path: Path) -> Path:
    """Create a simple CSV pricing file."""
    import pandas as pd
    path = tmp_path / "test_pricing.csv"
    df = pd.DataFrame({
        "trade": ["finishing", "mep", "civil"],
        "item": ["Floor tiles", "A/C split unit", "Concrete slab"],
        "unit": ["m2", "nr", "m3"],
        "rate": [85, 1200, 450],
    })
    df.to_csv(str(path), index=False)
    return path


SAMPLE_BOQ_TEXT = """
PROJECT: Office Fit-Out — Level 4
Total Area: 500 m2
Grade: Premium

BILL OF QUANTITIES
Item    Description                     Unit    Qty     Rate    Amount
1.01    Raised access floor (150mm)     m2      400     220     88000
1.02    Suspended ceiling grid          m2      400     130     52000
1.03    Glass partition system          lm      60      850     51000
1.04    Ceramic floor tiles 600x600     m2      350     85      29750
2.01    Split A/C units (2.5 ton)       nr      15      1200    18000
2.02    Electrical distribution board   nr      3       2500    7500
"""


# ─────────────────────────────────────────────
# PDF Parser Tests
# ─────────────────────────────────────────────

def test_pdf_parser_text_extraction(tmp_path):
    """PDF parser extracts text from a text-based PDF."""
    pdf_path = make_test_pdf(tmp_path)
    from parsers.pdf_parser import parse_pdf
    result = parse_pdf(str(pdf_path), job_id="TEST")

    assert result["document_type"] == "pdf"
    assert result["page_count"] >= 1
    assert isinstance(result["extracted_text"], str)
    assert isinstance(result["tables"], list)
    assert isinstance(result["confidence"], float)
    assert 0.0 <= result["confidence"] <= 1.0
    assert isinstance(result["flagged_gaps"], list)
    print(f"  PDF: pages={result['page_count']} chars={len(result['extracted_text'])} confidence={result['confidence']:.2f}")


def test_pdf_parser_missing_file():
    """PDF parser raises FileNotFoundError for missing files."""
    from parsers.pdf_parser import parse_pdf
    with pytest.raises(FileNotFoundError):
        parse_pdf("/nonexistent/path/file.pdf")


# ─────────────────────────────────────────────
# DOCX Parser Tests
# ─────────────────────────────────────────────

def test_docx_parser_extracts_text_and_tables(tmp_path):
    """DOCX parser extracts paragraphs and tables."""
    docx_path = make_test_docx(tmp_path)
    from parsers.docx_parser import parse_docx
    result = parse_docx(str(docx_path), job_id="TEST")

    assert result["document_type"] == "docx"
    assert "450" in result["extracted_text"] or "Office" in result["extracted_text"]
    assert len(result["tables"]) >= 1

    table = result["tables"][0]
    assert "headers" in table
    assert "rows" in table
    assert len(table["rows"]) >= 3
    print(f"  DOCX: tables={len(result['tables'])} rows={table['row_count']} confidence={result['confidence']:.2f}")


def test_docx_parser_missing_file():
    from parsers.docx_parser import parse_docx
    with pytest.raises(FileNotFoundError):
        parse_docx("/nonexistent/file.docx")


# ─────────────────────────────────────────────
# Excel Parser Tests
# ─────────────────────────────────────────────

def test_excel_parser_reads_boq(tmp_path):
    """Excel parser reads BOQ data from .xlsx."""
    xlsx_path = make_test_excel(tmp_path)
    from parsers.excel_parser import parse_excel
    result = parse_excel(str(xlsx_path), job_id="TEST")

    assert result["document_type"] in ("xlsx", "xls")
    assert len(result["tables"]) >= 1

    table = result["tables"][0]
    assert table["row_count"] == 5
    assert "Description" in table["headers"] or any("desc" in h.lower() for h in table["headers"])
    assert result["confidence"] >= 0.6
    print(f"  Excel: sheets={result['page_count']} rows={table['row_count']} confidence={result['confidence']:.2f}")


def test_csv_parser_reads_data(tmp_path):
    """Excel parser handles CSV files."""
    csv_path = make_test_csv(tmp_path)
    from parsers.excel_parser import parse_excel
    result = parse_excel(str(csv_path), job_id="TEST")

    assert result["document_type"] == "csv"
    assert result["tables"][0]["row_count"] == 3
    assert "trade" in result["tables"][0]["headers"]
    print(f"  CSV: rows={result['tables'][0]['row_count']} confidence={result['confidence']:.2f}")


def test_excel_parser_missing_file():
    from parsers.excel_parser import parse_excel
    with pytest.raises(FileNotFoundError):
        parse_excel("/nonexistent/file.xlsx")


# ─────────────────────────────────────────────
# Dispatcher Tests
# ─────────────────────────────────────────────

def test_dispatcher_routes_pdf(tmp_path):
    """Dispatcher correctly routes .pdf to pdf parser."""
    pdf_path = make_test_pdf(tmp_path)
    from parsers.dispatcher import dispatch
    result = dispatch(str(pdf_path), job_id="TEST")
    assert result["document_type"] == "pdf"
    print("  Dispatcher: PDF → pdf_parser ✓")


def test_dispatcher_routes_docx(tmp_path):
    """Dispatcher correctly routes .docx to docx parser."""
    docx_path = make_test_docx(tmp_path)
    from parsers.dispatcher import dispatch
    result = dispatch(str(docx_path), job_id="TEST")
    assert result["document_type"] == "docx"
    print("  Dispatcher: DOCX → docx_parser ✓")


def test_dispatcher_routes_excel(tmp_path):
    """Dispatcher correctly routes .xlsx to excel parser."""
    xlsx_path = make_test_excel(tmp_path)
    from parsers.dispatcher import dispatch
    result = dispatch(str(xlsx_path), job_id="TEST")
    assert result["document_type"] in ("xlsx", "xls")
    print("  Dispatcher: XLSX → excel_parser ✓")


def test_dispatcher_rejects_unsupported():
    """Dispatcher raises ValueError for unsupported file types."""
    from parsers.dispatcher import dispatch
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"test")
        tmp = f.name
    try:
        with pytest.raises(ValueError, match="Unsupported"):
            dispatch(tmp)
    finally:
        os.unlink(tmp)


def test_dispatcher_missing_file():
    """Dispatcher raises FileNotFoundError for missing files."""
    from parsers.dispatcher import dispatch
    with pytest.raises(FileNotFoundError):
        dispatch("/no/such/file.pdf")


def test_dispatch_many_continues_on_error(tmp_path):
    """dispatch_many returns partial results even if one file fails."""
    xlsx_path = make_test_excel(tmp_path)
    from parsers.dispatcher import dispatch_many
    results = dispatch_many([str(xlsx_path), "/nonexistent/file.pdf"], job_id="TEST")
    assert len(results) == 2
    assert results[0]["document_type"] in ("xlsx", "xls")
    assert results[1]["document_type"] == "error"
    assert results[1]["confidence"] == 0.0
    print("  dispatch_many: continues on error ✓")


# ─────────────────────────────────────────────
# Agent #1 Unit Tests (no LLM)
# ─────────────────────────────────────────────

def test_agent01_merge_documents():
    """Agent #1 _merge_documents correctly combines multiple docs."""
    from agents.agent_01_ingestion import _merge_documents, _build_error_doc
    from core.state import DocumentInfo

    docs: list[DocumentInfo] = [
        {
            "file_path": "boq.xlsx",
            "document_type": "xlsx",
            "raw_text": "BOQ content here",
            "tables": [{"headers": ["Item"], "rows": [{"Item": "1.01"}], "row_count": 1}],
            "dimensions": [{"element": "room_1", "length": 10, "width": 5}],
            "material_specs": [{"trade": "finishing", "item": "tiles"}],
            "floor_areas": [{"floor_name": "Ground", "area_m2": 350}],
            "boq_items": [],
            "scope_statements": ["All finishing works"],
            "images_extracted": [],
            "confidence": 0.85,
            "flagged_gaps": [],
        },
        {
            "file_path": "spec.docx",
            "document_type": "docx",
            "raw_text": "Spec content here",
            "tables": [],
            "dimensions": [],
            "material_specs": [{"trade": "mep", "item": "AC units"}],
            "floor_areas": [],
            "boq_items": [],
            "scope_statements": ["MEP works included"],
            "images_extracted": [],
            "confidence": 0.75,
            "flagged_gaps": ["No tables found"],
        },
    ]

    from core.logger import get_logger
    log = get_logger("test", "TEST")
    merged = _merge_documents(docs, log)

    assert len(merged["documents"]) == 2
    assert len(merged["all_tables"]) == 1
    assert len(merged["all_dimensions"]) == 1
    assert len(merged["all_material_specs"]) == 2
    assert len(merged["all_floor_areas"]) == 1
    assert merged["overall_confidence"] == pytest.approx(0.80, abs=0.01)
    assert "BOQ content here" in merged["combined_text"]
    assert "Spec content here" in merged["combined_text"]
    print(f"  Merge OK: confidence={merged['overall_confidence']:.2f} gaps={merged['flagged_gaps']}")


def test_agent01_empty_file_list():
    """Agent #1 handles empty file list gracefully."""
    from core.state import EstimatorState
    from agents.agent_01_ingestion import run

    state: EstimatorState = {
        "job_id": "TEST001",
        "status": "ingestion",
        "user_description": "Office fit-out",
        "uploaded_file_paths": [],
        "project_context": {"project_type": "office_fitout", "grade": "premium"},
        "intake_confirmed": True,
        "checkpoint_1_approved": False,
        "checkpoint_2_approved": False,
        "checkpoint_3_approved": False,
        "errors": [],
        "warnings": [],
        "agent_timings": {},
    }

    result = run(state)
    assert result["extracted_data"] is not None
    assert result["extracted_data"]["overall_confidence"] == 0.0
    assert len(result["warnings"]) > 0
    print(f"  Empty file list handled: warning='{result['warnings'][0]}'")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
