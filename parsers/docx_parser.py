"""
parsers/docx_parser.py
Parses .docx files using python-docx.
Extracts paragraphs, tables, and document metadata.
"""

from __future__ import annotations
from pathlib import Path
from core.logger import get_logger


def parse_docx(file_path: str, job_id: str = "system") -> dict:
    """
    Parse a .docx file and return structured content.

    Returns:
        {
            "filename": str,
            "document_type": "docx",
            "page_count": int (estimated),
            "extracted_text": str,
            "tables": [...],
            "is_scanned": False,
            "confidence": float,
            "flagged_gaps": [...]
        }
    """
    from docx import Document

    log_ctx = get_logger("docx_parser", job_id)
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    log_ctx.info(f"Parsing DOCX: {path.name}")
    doc = Document(str(path))

    # ── Extract paragraphs ───────────────────────────────────────
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    # ── Extract tables ───────────────────────────────────────────
    tables = []
    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # Use first row as headers
        headers = [cell.text.strip() or f"col_{i}" for i, cell in enumerate(table.rows[0].cells)]
        rows = []
        for row in table.rows[1:]:
            row_dict = {}
            for i, cell in enumerate(row.cells):
                key = headers[i] if i < len(headers) else f"col_{i}"
                row_dict[key] = cell.text.strip()
            rows.append(row_dict)

        tables.append({
            "page": None,  # DOCX doesn't have explicit page numbers
            "table_index": t_idx,
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
        })

    # Estimate page count (~500 words per page)
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 500)

    flagged_gaps = []
    if len(full_text.strip()) < 100:
        flagged_gaps.append("Very little text extracted from DOCX")
    if not tables:
        flagged_gaps.append("No tables found in DOCX")

    confidence = _score_confidence(full_text, tables)

    log_ctx.info(
        f"DOCX parsed: est_pages={estimated_pages} tables={len(tables)} "
        f"chars={len(full_text)} confidence={confidence:.2f}"
    )

    return {
        "filename": path.name,
        "document_type": "docx",
        "page_count": estimated_pages,
        "extracted_text": full_text,
        "tables": tables,
        "is_scanned": False,
        "confidence": confidence,
        "flagged_gaps": flagged_gaps,
    }


def _score_confidence(text: str, tables: list) -> float:
    score = 0.0
    if len(text) > 500:
        score += 0.5
    elif len(text) > 100:
        score += 0.3
    if tables:
        score += 0.3
    if len(text) > 2000:
        score += 0.2
    return min(score, 1.0)
